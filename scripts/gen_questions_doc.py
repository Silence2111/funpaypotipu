#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация Word-документа с открытыми развилками и вопросами по проекту GameMarket (fp-pl)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# базовый шрифт
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=10)
        shade_cell(hdr[i], "1F3864")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def h_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)

def para(text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

def answer_line(label="Решение / комментарий:"):
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = GREEN
    r2 = p.add_run("_" * 70)
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

def reco(text):
    p = doc.add_paragraph()
    r = p.add_run("Рекомендация: ")
    r.bold = True
    r.font.color.rgb = GREEN
    p.add_run(text)

# ---------------- Титул ----------------
h_title("GameMarket (fp-pl)")
para("Открытые развилки и вопросы для решения", color=GREY, size=14)
meta = doc.add_paragraph()
meta.add_run("Дата: 17.06.2026   •   Статус: Фаза 0 (проектирование)   •   "
             "Проект: маркетплейс цифровых игровых товаров (FunPay/Playerok + Kupikod/GGSel)").font.size = Pt(9)
meta.runs[0].font.color.rgb = GREY
doc.add_paragraph()

para("Документ фиксирует решения, уже принятые по архитектуре, и собирает открытые "
     "вопросы, которые требуют выбора со стороны бизнеса/владельца. По каждой развилке "
     "приведены варианты, плюсы/минусы и рекомендация. Поля «Решение / комментарий» — "
     "для фиксации финального выбора.", italic=True, color=GREY)

# ---------------- Зафиксированные решения ----------------
h1("1. Уже зафиксированные решения")
add_table(
    ["Тема", "Решение"],
    [
        ["Стек", "TypeScript-монорепо (pnpm + Turborepo): NestJS (API) + Next.js (фронт, SSR/SEO) + PostgreSQL/Prisma + Redis/BullMQ + Socket.IO + S3/MinIO"],
        ["Модель продукта", "Гибрид на одном ядре: P2P-эскроу (как FunPay/Playerok) + авто-выдача/пополнения (как Kupikod/GGSel) как стратегии выдачи одного заказа"],
        ["Деньги", "Двойная бухгалтерия (ledger), append-only, идемпотентность, эскроу-счёт — самый строгий модуль"],
        ["SEO", "SSR/ISR всех публичных страниц — органика главный канал лидеров рынка"],
        ["Первый шаг", "Сначала максимальная плановая архитектура (12 документов в docs/), код — следующим этапом"],
    ],
    widths=[1.6, 4.9],
)

# ---------------- Открытые развилки ----------------
h1("2. Открытые развилки (нужно решение)")

# Развилка 1
h2("2.1. Модель комиссии")
para("Сколько и с кого берём комиссию со сделки — главный рычаг дохода.")
add_table(
    ["Вариант", "Плюсы", "Минусы"],
    [
        ["Одинарная (только с продавца) — как FunPay", "Привычно продавцам, проще цена для покупателя", "Ниже доход с одной сделки"],
        ["Двойная (с продавца и покупателя) — как Playerok", "Выше доход, гибкость", "Цена «кусается», возможен отток к конкурентам"],
    ],
    widths=[2.3, 2.2, 2.0],
)
reco("Архитектура (FeeRule) поддерживает обе и разные ставки по категориям. Нужен дефолт на старте.")
answer_line()
doc.add_paragraph()

# Развилка 2
h2("2.2. Платёжный провайдер (ввод/вывод денег)")
para("Рынок «серый» и подвижный. Определяет, какую первую реализацию PaymentProvider пишем в Фазе 2.")
add_table(
    ["Вариант", "Комментарий"],
    [
        ["СБП / карты (ЮKassa и аналоги)", "Массовость, удобство, но требует юрлица и комплаенса"],
        ["Криптовалюта (USDT и пр.)", "Обходит ограничения, популярно в нише, выше порог входа для части аудитории"],
        ["Card-to-card / P2P-переводы", "Гибко, но выше риск и ручной труд"],
        ["Несколько сразу", "Шире охват, но дольше старт"],
    ],
    widths=[2.4, 4.1],
)
reco("Стартовать с одного провайдера через абстракцию, остальные добавлять без изменения ядра.")
answer_line()
doc.add_paragraph()

# Развилка 3
h2("2.3. Юрисдикция и легал")
para("Влияет на KYC, налоги, договор-оферту, доступ к платёжным провайдерам. "
     "Playerok вынесен в Казахстан не случайно.")
add_table(
    ["Вариант", "Комментарий"],
    [
        ["РФ (юрлицо/ИП)", "Проще локальные платежи, но регуляторные риски по «серому» ассортименту"],
        ["Казахстан / др. СНГ", "Модель Playerok: мягче регулирование, доступ к зарубежным платежам"],
        ["Зарубежная юрисдикция", "Гибкость по платежам/крипте, сложнее с локальными методами и поддержкой"],
    ],
    widths=[2.0, 4.5],
)
reco("Решение бизнес-уровня; архитектура нейтральна (KYC-уровни и оферта конфигурируемы).")
answer_line()
doc.add_paragraph()

# Развилка 4
h2("2.4. Стартовая вертикаль (ниша входа)")
para("Бить во всё сразу размазывает усилия и ликвидность. Отчёт: окно входа — не лобовая "
     "конкуренция с FunPay, а узкая вертикаль / Telegram-нативность / B2B.")
add_table(
    ["Вариант", "Комментарий"],
    [
        ["Аккаунты + валюта 1–2 популярных игр", "Быстро набрать ликвидность в узкой нише, потом расширять"],
        ["Пополнения/донат (высокая маржа)", "Сегмент Kupikod/GGSel, но завязан на внешние интеграции"],
        ["Telegram-нативная площадка", "Дифференциатор по дистрибуции, ниже стоимость входа"],
        ["Широкий охват сразу", "Дольше до ликвидности, выше риск распыления"],
    ],
    widths=[2.6, 3.9],
)
reco("Выбрать 1 узкую вертикаль для набора ликвидности; архитектура уже расширяема на все сегменты.")
answer_line()
doc.add_paragraph()

# Развилка 5
h2("2.5. Жёсткость антискама (маскирование контактов в чате)")
para("Баланс между защитой комиссии (не давать увести сделку мимо эскроу) и лояльностью честных продавцов.")
add_table(
    ["Вариант", "Плюсы", "Минусы"],
    [
        ["Жёстко (block) — не пропускать контакты", "Максимум защиты комиссии и покупателя", "Раздражает честных, трение в общении"],
        ["Мягко (mask/flag) — скрывать/помечать", "Меньше трения, сигнал в антифрод", "Часть сделок всё же уходит мимо"],
        ["Гибко по уровню доверия", "Послабления верифицированным/высокорейтинговым", "Сложнее настройка правил"],
    ],
    widths=[2.4, 2.1, 2.0],
)
reco("Настраиваемые уровни (system_setting) + послабления для проверенных. Нужен дефолт на старте.")
answer_line()
doc.add_paragraph()

# ---------------- Дополнительные вопросы ----------------
h1("3. Дополнительные вопросы (всплывут по ходу)")
para("Не блокируют старт, но их стоит зафиксировать заранее:")
for q in [
    "Провайдер пополнений Steam/донат — какой интегрируем первым (Фаза 5).",
    "Валюта старта (RUB?) и когда включать мультивалютность.",
    "Бренд: название, домен, логотип, тон коммуникации.",
    "Мобильное приложение и/или Telegram Mini App — приоритет и сроки.",
    "Гарантийные периоды и холды выплат по категориям (конкретные числа для антифрода).",
    "KYC-уровни: что и когда требуем у продавцов (для повышенных лимитов/выплат).",
    "Лимиты сумм/частоты операций по уровням доверия — конкретные пороги.",
    "Хостинг/инфраструктура: где разворачиваем (один VPS → managed → K8s).",
    "Команда, бюджет и желаемые сроки запуска MVP — влияет на темп фаз.",
    "Поддержка/арбитраж: кто и как разбирает споры на старте (ручной процесс).",
]:
    bullet(q)

# ---------------- Следующий шаг ----------------
h1("4. Следующий шаг (выбрать направление)")
add_table(
    ["Вариант", "Что делаем"],
    [
        ["A. Разобрать развилки", "Зафиксировать ответы по разделу 2 → обновить архитектурные документы"],
        ["B. Перейти к коду (Фаза 0)", "Установить Node.js + pnpm (сейчас в системе их нет), собрать каркас монорепо"],
        ["C. Углубить раздел", "Например, перенести доменную модель (docs/02) в реальную Prisma-схему"],
    ],
    widths=[2.2, 4.3],
)
para("Примечание: Node.js и pnpm в системе пока не установлены (есть Python 3.9 и git). "
     "Установка — первый технический шаг перед написанием кода.", italic=True, color=GREY, size=10)

out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "GameMarket_Развилки_и_вопросы.docx")
doc.save(out)
print("Сохранено:", out)
